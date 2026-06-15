# -*- coding: utf-8 -*-
"""
Generates the retail-expansion gap-analysis deliverables:
  - Retail_Expansion_Gap_Analysis.xlsx  (gap register + matrix + roadmap + data model)
  - Retail_Expansion_Gap_Analysis.docx  (narrative report)

Synthesised from a cross-repo exploration of the Flutter app (kirana_ai) and the
FastAPI backend (kirana-master-backend). Re-run to regenerate.
"""
from __future__ import annotations

import os
from datetime import date

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
TODAY = "2026-06-15"

# ── Scope ────────────────────────────────────────────────────────────────────
VERTICALS_IN = [
    "Apparel / Clothing", "Footwear", "Electronics / Mobile", "Stationery / Books",
    "Hardware / Paint", "Cosmetics / Beauty", "Toys / Gifts",
    "Home & Kitchenware", "General / Variety stores",
]
VERTICALS_OUT = ["Medical / Pharmacy (drug licence, scheduled drugs, prescriptions, mandatory batch+expiry — high regulatory burden, explicitly excluded)"]

# Matrix vertical columns (compact)
MX_COLS = ["Grocery (Kirana)", "Apparel", "Footwear", "Electronics", "Stationery", "Hardware", "Cosmetics", "General"]

# ── Gap register ─────────────────────────────────────────────────────────────
# id, theme, layer, current, gap_impact, change, verticals, priority, effort, evidence
GAPS = [
    ("G-01", "Vertical config", "Backend+FE+DB",
     "store.store_type exists (default 'kirana') but branches NO behaviour; app is grocery-first everywhere.",
     "There is no notion of a 'vertical' that drives features, units, KPIs, copy, ML or tax. Every vertical sees grocery UI/logic.",
     "Introduce a Vertical Config layer (per store/vertical): a feature-flag + config object that drives enabled features, unit registry, attribute schema, KPI set, ML profile, tax profile and copy pack. Single keystone that unlocks the rest.",
     "All", "P0", "XL",
     "ensure_full_schema.py:59 store_type; no 'WHERE store_type=' branching anywhere"),

    ("G-02", "Onboarding", "Frontend+Backend",
     "Onboarding store-type list is grocery-flavoured: kirana / general / provision / fruits_vegetables / pharmacy / stationery / other.",
     "Missing apparel, footwear, electronics, hardware, cosmetics; choice is cosmetic (stored, never used). No sub-vertical or attribute config captured.",
     "Replace with a real vertical picker that maps to a Vertical Config; capture sub-type + unit/attribute defaults; gate onboarding copy by vertical.",
     "All", "P0", "M",
     "onboarding/views/steps/business_step.dart:14-22"),

    ("G-03", "Vertical config", "Frontend",
     "Expiry, loose-weight, perishable and weight-variant UI are always shown.",
     "Non-grocery users see irrelevant fields (expiry pickers, 'sold by weight (e.g. Maida, Pulse)').",
     "Feature-flag every grocery-only widget; render by vertical (expiry/loose OFF for apparel/electronics/stationery/hardware; optional for cosmetics).",
     "All non-grocery", "P0", "M",
     "add_product_sheet.dart perishable/expiry; app_en.arb:919 invLooseItemSub"),

    ("G-04", "Product data model", "DB+Backend+FE",
     "No variant matrix. 'Variants' today are separate product rows differing by weight (1kg vs 500g) each with its own product_id.",
     "Apparel/footwear need ONE product with a grid of size x colour variants; today this is impossible without exploding the catalog and breaking reporting.",
     "Add product_variant table (variant_id, product_id, sku, barcode, attributes, price, mrp, cost, stock); inventory & orders reference variant; FE size x colour grid.",
     "Apparel, Footwear, Cosmetics, Electronics", "P0", "XL",
     "pos_product.dart; add_product_sheet.dart _VariantData (weight-based)"),

    ("G-05", "Product data model", "DB+Backend+FE",
     "Product attributes are a fixed grocery set (unit, weight, is_loose, is_perishable, barcode, mrp).",
     "No way to store size, colour, material, RAM/storage, shade, dimensions, model etc. per vertical.",
     "Add a dynamic attribute schema (product_attribute_def per vertical + values on product/variant). Drives add-product form and filters.",
     "All non-grocery", "P0", "L",
     "ensure_full_schema.py product table 111-127"),

    ("G-06", "Product data model", "DB+Backend+FE",
     "No serial / IMEI tracking (unit-level identity).",
     "Electronics/mobile must track each unit (serial/IMEI) for sale, warranty and returns.",
     "Add product_serial (serial/imei, status, warranty_until); POS sells a specific serial; receive/return by serial.",
     "Electronics", "P1", "L",
     "no serial column in product/inventory"),

    ("G-07", "Product data model", "DB+Backend+FE",
     "No warranty concept.",
     "Electronics, some hardware/appliances sell with warranty (months / tiers / RMA).",
     "Add warranty fields (warranty_months, warranty_until) + optional warranty-tier pricing; surface on bill + returns.",
     "Electronics, Hardware", "P1", "M",
     "no warranty fields anywhere"),

    ("G-08", "Inventory", "DB+Backend",
     "Perishable / batch / expiry is first-class (inventory_batch.expiry_date NOT NULL, markdown, waste).",
     "Non-grocery has no expiry; mandatory batch+expiry is wrong for apparel/electronics/stationery/hardware.",
     "Make batch/expiry optional & vertical-gated; per-variant stock; keep for grocery & (optional) cosmetics.",
     "All non-grocery", "P0", "M",
     "ensure_full_schema.py inventory_batch 657-671 (expiry_date NOT NULL)"),

    ("G-09", "Units of measure", "Frontend+Backend",
     "Units hardcoded: pcs, kg, g, L, ml, dozen, pack, box, bundle; loose auto-switches to kg.",
     "Apparel needs piece/pair/set; electronics unit/box; hardware m/ft/pack; the grocery list is noise elsewhere.",
     "Vertical-aware unit registry (backend-driven); default per vertical; remove grocery example copy.",
     "All", "P0", "M",
     "add_product_sheet.dart:17 _units; app_en.arb:919"),

    ("G-10", "Pricing & Tax", "DB+Backend+FE",
     "GST is parsed from supplier invoices but NEVER applied at pricing/POS; no tax bracket logic; pricing table has no tax column.",
     "Higher-value verticals need compliant tax: apparel 5%/12% by INR1000 threshold, footwear similar, electronics 18%/28%, grocery 0%/5%.",
     "Add tax engine: HSN/category + price-threshold -> GST rate; product.gst_rate/hsn; apply at checkout; show tax breakup on bill & receipt.",
     "All (compliance)", "P0", "L",
     "ai/routes.py invoice prompt parses GST; pricing table has no tax (ensure_full_schema.py 150-160)"),

    ("G-11", "Pricing & Tax", "DB+Backend",
     "Cost & price are per product, not per variant; cost lives in product_supplier.",
     "Apparel cost/price varies by size/colour; electronics by model/config.",
     "Move price/mrp/cost to variant level; keep product-level fallback.",
     "Apparel, Footwear, Electronics", "P1", "M",
     "pricing table product-level (ensure_full_schema.py 150-160)"),

    ("G-12", "Pricing & Tax", "Backend+FE",
     "Markdown/clearance is expiry-driven (sell before expiry).",
     "Apparel/footwear need season/end-of-line markdown engine (not expiry).",
     "Add season/markdown pricing (planned price drops, clearance flags) decoupled from expiry.",
     "Apparel, Footwear, Cosmetics", "P2", "M",
     "near_expiry markdown flow; inventory_batch.markdown_pct"),

    ("G-13", "KPIs & analytics", "Backend",
     "KPI registry is grocery-coded: perishable-waste, dead-stock@21d, category-mix, festive-uplift, daily-velocity assumptions.",
     "These mislead or N/A for other verticals; dead-stock@21d wrongly flags healthy slow-movers (denim, laptops).",
     "Vertical KPI sets selected by Vertical Config; per-vertical thresholds/windows.",
     "All non-grocery", "P1", "L",
     "kpis/registry.py; calculator.py calc_perishable_waste / calc_dead_stock"),

    ("G-14", "KPIs & analytics", "Backend",
     "No apparel KPIs.",
     "Apparel needs sell-through %, size-curve / size availability, season & markdown %, GMROI.",
     "Implement apparel KPI pack.",
     "Apparel, Footwear", "P1", "M", "missing"),

    ("G-15", "KPIs & analytics", "Backend",
     "No electronics KPIs.",
     "Electronics needs attach-rate, warranty-claim %, GMROI, model lifecycle, serial ageing.",
     "Implement electronics KPI pack.",
     "Electronics", "P1", "M", "missing"),

    ("G-16", "KPIs & analytics", "Frontend+Backend",
     "Dashboard / intelligence copy is grocery-shaped ('critical SKUs', 'reorder today', stockout-risk strip).",
     "Wrong tone/metrics for other verticals.",
     "Vertical-aware dashboard copy & recommendation types.",
     "All non-grocery", "P1", "M",
     "app_en.arb dashBriefing*/dashMetric*; recommendation_model.dart"),

    ("G-17", "ML / forecasting", "ML",
     "Demand/stockout/deadstock/velocity models tuned for grocery daily velocity & perishability.",
     "Apparel is seasonal & size-curve driven; electronics slow & high-value; grocery thresholds misfire.",
     "Per-vertical ML profiles (features, windows, thresholds); seasonal models for apparel; cold-start heuristics for new stores.",
     "All non-grocery", "P1", "XL",
     "ml_models/models/deadstock_detector.py DEAD_STOCK_DAYS=21; sku_velocity percentiles"),

    ("G-18", "ML / forecasting", "ML+Backend",
     "Models are global, not vertical-scoped.",
     "Cannot load apparel/electronics model variants.",
     "Vertical-scoped model artefacts + config; selection via Vertical Config.",
     "All non-grocery", "P1", "M", "ml_models config global"),

    ("G-19", "AI intake (voice/handwrite)", "Backend",
     "Voice & handwrite Gemini prompts hardcode grocery dictionaries (rice/dal/atta/oil in TE/HI/...).",
     "No grounding for jeans/saree/shirt/size-32 or mobile/laptop/warranty.",
     "Vertical-aware prompts + ground against the store's own catalog/variants instead of a fixed grocery glossary.",
     "All non-grocery", "P1", "M",
     "ai/routes.py:299-372 _VOICE_PROMPT/_HANDWRITE_PROMPT"),

    ("G-20", "AI intake (invoice)", "Backend",
     "Invoice OCR prompt is grocery-flavoured (MRP-on-packet, snack examples).",
     "Apparel/electronics invoices carry size/colour/serial/warranty per line.",
     "Vertical invoice prompt variants; map to variants/serials.",
     "Apparel, Electronics", "P1", "M",
     "ai/routes.py:374-441 _INVOICE_PROMPT"),

    ("G-21", "Vision", "Backend+FE",
     "Vision shelf-scan + Gemini prompt are grocery-specific ('kirana shelf', readable package text, morning/evening cadence).",
     "Apparel is folded/hung with no readable text; electronics are boxed/barcoded; shelf-scan won't recognise them.",
     "Gate Vision by vertical (off where unsupported); medium-term redesign capture per vertical (rack/box/barcode); retrain per vertical.",
     "Apparel, Electronics, others", "P1", "L",
     "vision/analyzer.py SHELF_PROMPT 'kirana store shelf'; vision_provider.dart shelf/analyze"),

    ("G-22", "POS & billing", "Frontend",
     "POS add-to-cart assumes loose weight dialog (decimal kg) or single discrete unit.",
     "Non-grocery needs variant selection (size/colour, model, serial) before add; no weighing.",
     "Variant-aware add-to-cart (pick size/colour/serial); hide weight dialog where not loose.",
     "Apparel, Footwear, Electronics", "P0", "M",
     "pos_tab.dart:66-130 _showWeightDialog"),

    ("G-23", "POS & billing", "Frontend+Backend",
     "Bill / thermal receipt has no tax breakup.",
     "Higher-value verticals & GST compliance need tax on the invoice/receipt.",
     "Add tax lines to bill + printer template (ties to G-10).",
     "All", "P1", "M",
     "printer receipt template; pos order"),

    ("G-24", "Copy & branding", "Frontend (l10n)",
     "'Kirana AI' brand + grocery examples (Maida/Pulse), 'mykiranastore' hints, grocery welcome/consent copy across 7 ARBs.",
     "Apparel/electronics users see grocery branding & nonsensical examples.",
     "Neutral / white-label or vertical-aware copy pack; parameterise brand & examples; re-translate the delta across 7 languages.",
     "All non-grocery", "P1", "L",
     "app_en.arb:22,49,132,178,919,157"),

    ("G-25", "WhatsApp", "Backend",
     "WhatsApp menu buttons/templates hardcode grocery KPIs (Stockout, Fast-Moving SKUs) per language.",
     "Apparel/electronics owners want different metrics & flows.",
     "Vertical-aware WhatsApp menus/templates selected by Vertical Config.",
     "All non-grocery", "P2", "M",
     "whatsapp/templates.py:48-66 ANALYTICS_BUTTONS"),

    ("G-26", "Catalog", "Backend+Data",
     "Catalog = ~11k grocery blinkit SKUs (KAI-{blinkitId}); add-product leans on it.",
     "Other verticals have no seed catalog; manual add is slow without barcode/catalog support.",
     "Barcode lookup (global) + per-vertical catalog ingestion + fast manual add with attribute templates.",
     "All non-grocery", "P1", "L",
     "db_generation/seed_blinkit_catalog.py; vision/matcher.py loads all product"),

    ("G-27", "Catalog", "Backend",
     "Catalog matcher loads ALL products in-memory; no vertical partition.",
     "Apparel/electronics with variants reach 50k+ rows; in-memory match doesn't scale & mixes verticals.",
     "Partition catalog by vertical; move to indexed / vector search for large catalogs.",
     "Apparel, Electronics", "P2", "L",
     "vision/matcher.py:21-93 SELECT all product"),

    ("G-28", "Procurement", "Backend+FE",
     "Supplier/PO model is generic but date-based; no MOQ, lead time, seasonal/size-pack buying.",
     "Apparel buys by season/size-pack; electronics receive by serial; reorder logic differs.",
     "Extend procurement with MOQ, lead time, season; serial receipt for electronics.",
     "Apparel, Electronics, Hardware", "P2", "M",
     "procurement_models.dart; purchases table"),

    ("G-29", "Customer", "DB+Backend",
     "customer.household_size (default 4) and grocery referral assumptions.",
     "Irrelevant for apparel/electronics; harmless but noisy.",
     "Make demographic fields optional / vertical-aware; keep generic segments.",
     "All non-grocery", "P2", "S",
     "ensure_full_schema.py customer.household_size"),

    ("G-30", "Credit (udhaar)", "Backend+FE",
     "Udhaar/khata credit ledger (KEEP — broadly used across Indian small retail).",
     "NOT a blocker: apparel, hardware, general stores also extend credit. Electronics lean to EMI.",
     "Keep udhaar as a cross-vertical strength; add optional EMI/instalment stub for electronics later.",
     "Electronics (EMI optional)", "P2", "S",
     "finance/khata — cross-vertical strength, not a gap"),

    ("G-31", "Notifications / widget", "Frontend",
     "Home widget + push taxonomy surface grocery KPIs (stockout, reorder, fast-moving).",
     "Wrong metrics for other verticals.",
     "Vertical-aware widget stats & notification copy.",
     "All non-grocery", "P2", "S",
     "home widget snapshot keys; notification taxonomy"),

    ("G-32", "Baskets / bundles", "Backend+FE",
     "Bundles/baskets are generic but use cases & copy are grocery (food combos).",
     "Minor: works for apparel ('outfit bundles'), needs copy.",
     "Vertical copy for bundles; otherwise reuse as-is (strength).",
     "All", "P2", "S",
     "features/baskets — generic engine"),

    ("G-33", "Reporting / export", "Backend+FE",
     "Reports/exports assume grocery categories & units.",
     "Need variant- and tax-aware reporting (size-curve, sell-through, tax summaries).",
     "Variant/tax-aware reports & exports per vertical.",
     "All non-grocery", "P2", "M",
     "transaction history; KPI exports"),

    ("G-34", "Data migration / rollout", "Backend+DB",
     "Existing grocery stores run on the current single-product model.",
     "Introducing variants/attributes/tax must not break live grocery stores or order history.",
     "Backward-compatible migration: variant is optional (grocery = single implicit variant); phase behind Vertical Config; verify on staging before deploy.",
     "All", "P0", "M",
     "single unified DB kirana_oltp; live grocery stores"),
]

# ── Feature x Vertical matrix ────────────────────────────────────────────────
# Works / Adapt / Build / Optional / N/A
MATRIX = [
    ("Orders & checkout core",        ["Works","Works","Works","Works","Works","Works","Works","Works"]),
    ("Product variants (size x colour)",["N/A","Build","Build","Adapt","Optional","Optional","Adapt","Optional"]),
    ("Serial / IMEI tracking",        ["N/A","N/A","N/A","Build","N/A","Optional","N/A","N/A"]),
    ("Expiry / batch tracking",       ["Works","N/A","N/A","N/A","N/A","N/A","Optional","Optional"]),
    ("Loose / weight POS",            ["Works","N/A","N/A","N/A","N/A","Optional","N/A","Optional"]),
    ("Units of measure",              ["Works","Adapt","Adapt","Adapt","Adapt","Adapt","Adapt","Adapt"]),
    ("Tax / GST engine",              ["Build","Build","Build","Build","Build","Build","Build","Build"]),
    ("KPIs / analytics",              ["Works","Build","Build","Build","Adapt","Adapt","Adapt","Adapt"]),
    ("Demand / ML forecast",          ["Works","Build","Build","Adapt","Adapt","Adapt","Adapt","Adapt"]),
    ("Voice / handwrite intake",      ["Works","Adapt","Adapt","Adapt","Adapt","Adapt","Adapt","Adapt"]),
    ("Vision shelf-scan",             ["Works","Build","Adapt","N/A","Adapt","N/A","Adapt","Adapt"]),
    ("WhatsApp assistant",            ["Works","Adapt","Adapt","Adapt","Adapt","Adapt","Adapt","Adapt"]),
    ("Udhaar / credit ledger",        ["Works","Works","Works","Optional","Works","Works","Works","Works"]),
    ("Baskets / bundles",             ["Works","Adapt","Adapt","Adapt","Works","Works","Adapt","Works"]),
    ("Catalog seed",                  ["Works","Build","Build","Build","Build","Build","Build","Adapt"]),
    ("Customer mgmt / segments",      ["Works","Works","Works","Works","Works","Works","Works","Works"]),
    ("Procurement / suppliers",       ["Works","Adapt","Adapt","Adapt","Works","Adapt","Works","Works"]),
    ("Customer-specific pricing",     ["Works","Works","Works","Works","Works","Works","Works","Works"]),
]

# ── Roadmap ──────────────────────────────────────────────────────────────────
ROADMAP = [
    ("Phase 0 — Foundation (vertical-enable the platform)", "~4-6 weeks",
     "Vertical Config layer + feature flags; dynamic attribute schema + product_variant table (DB+API+core add-product UI); vertical unit registry; tax engine (HSN/threshold -> GST, applied at POS + receipt); neutralise branding/copy scaffolding; backward-compatible migration (grocery = single implicit variant).",
     "Unlocks General/Variety, Stationery, Hardware, Cosmetics with light extra work.",
     "G-01, G-02, G-03, G-04, G-05, G-08, G-09, G-10, G-22, G-24, G-34"),
    ("Phase 1 — Apparel pilot", "~6-8 weeks",
     "Size x colour variant grid UI; apparel KPI pack (sell-through, size-curve, GMROI, season/markdown); apparel ML profile (seasonal) + cold-start heuristics; vertical copy pack; vision gated off; pilot with a few stores.",
     "First true non-grocery vertical live; validates the model.",
     "G-04, G-11, G-13, G-14, G-16, G-17, G-19, G-21, G-24"),
    ("Phase 2 — Electronics / Mobile", "~6-8 weeks",
     "Serial/IMEI + warranty; GST 18/28; electronics KPI pack (attach-rate, GMROI, warranty-claim); EMI/instalment stub; invoice prompt v2 (serial/warranty).",
     "High-value vertical with unit-level identity.",
     "G-06, G-07, G-10, G-15, G-20, G-30"),
    ("Phase 3 — Scale & polish", "Ongoing",
     "Per-vertical catalogs + ingestion; catalog partition / vector search; vertical AI prompts & WhatsApp menus; vision redesign per vertical; ML retraining; vertical-aware widget/notifications; reporting/exports.",
     "Breadth across the remaining verticals + quality.",
     "G-12, G-18, G-21, G-25, G-26, G-27, G-28, G-29, G-31, G-32, G-33"),
]

# ── Proposed data model ──────────────────────────────────────────────────────
DATAMODEL = [
    ("vertical_config", "New",
     "Per vertical (and/or per store): feature flags + config that drive the whole app.",
     "vertical_code, features(jsonb: expiry,loose,variants,serial,warranty...), unit_set, attribute_set, kpi_set, ml_profile, tax_profile, copy_pack"),
    ("product_attribute_def", "New",
     "Dynamic attribute definitions per vertical (drives add-product form + filters).",
     "vertical_code, attr_code, label, type(text|enum|number), options, is_variant_axis, sort"),
    ("product_variant", "New (keystone)",
     "One product -> many sellable variants (size x colour, model/config, shade). Grocery = single implicit variant.",
     "variant_id PK, product_id FK, sku, barcode, attributes(jsonb), price, mrp, cost, is_active"),
    ("product_serial", "New",
     "Unit-level identity for electronics (and serialised hardware).",
     "serial_id PK, variant_id FK, serial_no/imei UNIQUE, status(in_stock|sold|returned), warranty_until, sold_order_id"),
    ("tax_rule", "New",
     "Category/HSN + price-threshold -> GST rate; applied at checkout.",
     "hsn/category, min_price, max_price, gst_rate; (product gains hsn_code, gst_rate override)"),
    ("inventory (modify)", "Modify",
     "Stock at VARIANT level; batch/expiry optional & vertical-gated.",
     "add variant_id; make batch/expiry nullable; keep product-level for grocery"),
    ("pricing (modify)", "Modify",
     "Price/mrp/cost can live at variant level; add tax columns.",
     "variant_id (nullable), gst_rate/hsn"),
    ("product (modify)", "Modify",
     "Grocery-only columns become optional & vertical-gated.",
     "is_perishable/is_loose/weight nullable; add hsn_code, gst_rate, brand/season as needed"),
    ("store (modify)", "Modify",
     "Activate the dormant store_type into a real vertical_code linked to vertical_config.",
     "vertical_code FK -> vertical_config"),
]

# ── Styling helpers ──────────────────────────────────────────────────────────
NAVY = "1F3864"; BLUE = "2E5496"; LIGHT = "DCE6F1"; GREY = "F2F2F2"
PRIO = {"P0": "FFC7CE", "P1": "FFEB9C", "P2": "C6EFCE"}
MX = {"Works": "C6EFCE", "Adapt": "FFEB9C", "Build": "FFC7CE", "Optional": "DDEBF7", "N/A": "F2F2F2"}
THIN = Side(style="thin", color="BFBFBF")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def _hdr(cell, text, fill=BLUE, color="FFFFFF", size=11):
    cell.value = text
    cell.font = Font(bold=True, color=color, size=size)
    cell.fill = PatternFill("solid", fgColor=fill)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = BORDER


def _cell(cell, text, wrap=True, fill=None, bold=False, align="left"):
    cell.value = text
    cell.alignment = Alignment(horizontal=align, vertical="top", wrap_text=wrap)
    cell.font = Font(bold=bold, size=10)
    if fill:
        cell.fill = PatternFill("solid", fgColor=fill)
    cell.border = BORDER


def build_excel(path):
    wb = Workbook()

    # Sheet 1: Executive Summary
    ws = wb.active
    ws.title = "Executive Summary"
    ws.sheet_view.showGridLines = False
    ws["A1"] = "Retail Expansion — Gap Analysis"
    ws["A1"].font = Font(bold=True, size=18, color=NAVY)
    ws["A2"] = f"Kirana AI -> multi-vertical retail platform  |  {TODAY}"
    ws["A2"].font = Font(italic=True, size=11, color="595959")
    rows = [
        ("", ""),
        ("Objective", "Expand the Kirana AI app + backend from Indian grocery (kirana) to all other retail verticals, EXCEPT medical/pharmacy."),
        ("Verticals IN", ", ".join(VERTICALS_IN)),
        ("Vertical OUT", VERTICALS_OUT[0]),
        ("", ""),
        ("Headline", "The platform is ~80% architecturally generic (auth, orders, customers, procurement, payments, baskets, customer pricing, credit/udhaar) but the product/inventory model, catalog, KPIs, ML, AI prompts, Vision and copy are grocery-coded."),
        ("Keystone gap", "There is no 'vertical' concept driving behaviour. store_type exists but is dormant. Everything else hangs off building a Vertical Config layer + a product variant/attribute model."),
        ("Biggest blockers (P0)", "Vertical config; product variant matrix (size x colour); dynamic attributes; make expiry/loose/weight optional; vertical unit registry; tax/GST engine applied at POS; variant-aware POS; backward-compatible migration."),
        ("Recommended approach", "Config-driven 'vertical packs' (one codebase), NOT per-vertical forks. One vertical dimension drives features, units, attributes, KPIs, ML profile, AI prompts, tax and copy."),
        ("Phasing", "P0 Foundation (~4-6 wks) -> P1 Apparel pilot (~6-8 wks) -> P2 Electronics (~6-8 wks) -> P3 scale & polish (ongoing)."),
        ("Counts", f"{len(GAPS)} gaps logged across {len(set(g[1] for g in GAPS))} themes. See 'Gap Register'. Feature coverage in 'Feature x Vertical Matrix'. Plan in 'Phased Roadmap'. Schema in 'Proposed Data Model'."),
        ("Why pharmacy is excluded", "Drug licensing, scheduled drugs, prescription workflows and mandatory batch+expiry create a heavy, separate regulatory product — out of scope by design."),
    ]
    r = 4
    for label, val in rows:
        ws.cell(r, 1, label).font = Font(bold=True, color=NAVY, size=11)
        ws.cell(r, 1).alignment = Alignment(vertical="top")
        c = ws.cell(r, 2, val)
        c.alignment = Alignment(wrap_text=True, vertical="top")
        c.font = Font(size=10)
        r += 1
    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 120
    # legend
    ws.cell(r + 1, 1, "Priority legend").font = Font(bold=True, color=NAVY)
    for i, (k, v) in enumerate({"P0 = foundational blocker": "FFC7CE", "P1 = needed for first non-grocery verticals": "FFEB9C", "P2 = scale / polish": "C6EFCE"}.items()):
        cc = ws.cell(r + 2 + i, 1, k)
        cc.fill = PatternFill("solid", fgColor=v)
        cc.font = Font(size=10)

    # Sheet 2: Gap Register
    ws = wb.create_sheet("Gap Register")
    ws.sheet_view.showGridLines = False
    headers = ["ID", "Theme", "Layer", "Current (grocery-specific) state", "Gap / impact",
               "Required change", "Verticals affected", "Priority", "Effort", "Evidence (file refs)"]
    widths = [7, 18, 16, 40, 40, 46, 20, 9, 8, 34]
    for c, (h, w) in enumerate(zip(headers, widths), start=1):
        _hdr(ws.cell(1, c), h)
        ws.column_dimensions[get_column_letter(c)].width = w
    for ri, g in enumerate(GAPS, start=2):
        for ci, val in enumerate(g, start=1):
            fill = PRIO.get(val) if ci == 8 else (GREY if ri % 2 == 0 else None)
            _cell(ws.cell(ri, ci), val, fill=fill,
                  bold=(ci == 1), align="center" if ci in (1, 8, 9) else "left")
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(GAPS)+1}"

    # Sheet 3: Feature x Vertical matrix
    ws = wb.create_sheet("Feature x Vertical Matrix")
    ws.sheet_view.showGridLines = False
    _hdr(ws.cell(1, 1), "Feature / capability")
    ws.column_dimensions["A"].width = 32
    for c, name in enumerate(MX_COLS, start=2):
        _hdr(ws.cell(1, c), name)
        ws.column_dimensions[get_column_letter(c)].width = 14
    for ri, (feat, vals) in enumerate(MATRIX, start=2):
        _cell(ws.cell(ri, 1), feat, bold=True)
        for ci, v in enumerate(vals, start=2):
            _cell(ws.cell(ri, ci), v, fill=MX.get(v, None), align="center")
    ws.freeze_panes = "B2"
    lr = len(MATRIX) + 3
    ws.cell(lr, 1, "Legend:").font = Font(bold=True)
    for i, (k, v) in enumerate(MX.items()):
        cc = ws.cell(lr, 2 + i, k); cc.fill = PatternFill("solid", fgColor=v)
        cc.alignment = Alignment(horizontal="center"); cc.font = Font(size=9)

    # Sheet 4: Roadmap
    ws = wb.create_sheet("Phased Roadmap")
    ws.sheet_view.showGridLines = False
    headers = ["Phase", "Timeline", "Scope / workstreams", "Outcome", "Gap IDs"]
    widths = [34, 14, 70, 38, 30]
    for c, (h, w) in enumerate(zip(headers, widths), start=1):
        _hdr(ws.cell(1, c), h)
        ws.column_dimensions[get_column_letter(c)].width = w
    for ri, row in enumerate(ROADMAP, start=2):
        for ci, val in enumerate(row, start=1):
            _cell(ws.cell(ri, ci), val, bold=(ci == 1))
    ws.freeze_panes = "A2"

    # Sheet 5: Proposed data model
    ws = wb.create_sheet("Proposed Data Model")
    ws.sheet_view.showGridLines = False
    headers = ["Table", "Change", "Purpose", "Key columns"]
    widths = [26, 16, 52, 60]
    for c, (h, w) in enumerate(zip(headers, widths), start=1):
        _hdr(ws.cell(1, c), h)
        ws.column_dimensions[get_column_letter(c)].width = w
    for ri, row in enumerate(DATAMODEL, start=2):
        for ci, val in enumerate(row, start=1):
            _cell(ws.cell(ri, ci), val, bold=(ci == 1))
    ws.freeze_panes = "A2"

    wb.save(path)
    return path


# ── Word ─────────────────────────────────────────────────────────────────────
def build_word(path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        return h

    def para(text, bold=False, italic=False, size=10.5):
        p = doc.add_paragraph()
        run = p.add_run(text); run.bold = bold; run.italic = italic
        run.font.size = Pt(size)
        return p

    def bullets(items):
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = h
            for p in c.paragraphs:
                for run in p.runs:
                    run.bold = True; run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            c._tc.get_or_add_tcPr()
        # header shading
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        for i in range(len(headers)):
            tcPr = t.rows[0].cells[i]._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "2E5496")
            tcPr.append(shd)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
                for p in cells[i].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        return t

    # Title
    title = doc.add_heading("Retail Expansion — Gap Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    para("From Kirana AI (grocery) to a multi-vertical retail platform", italic=True, size=12)
    para(f"Prepared: {TODAY}   |   Scope: Flutter app (kirana_ai) + FastAPI backend (kirana-master-backend)", italic=True, size=9.5)

    # 1. Executive summary
    heading("1. Executive summary", 1)
    para("The company wants to take the existing Kirana AI app and backend — today purpose-built for Indian grocery (kirana) stores — and make it work for all other retail verticals (clothing, footwear, electronics, stationery, hardware, cosmetics, toys, home & kitchenware, general/variety stores), explicitly EXCLUDING medical/pharmacy.")
    para("Good news: the platform is roughly 80% architecturally generic. Authentication, orders/checkout, customers & segments, procurement, payments, baskets/bundles, customer-specific pricing and the udhaar (credit) ledger are vertical-neutral and reusable as-is.")
    para("The constraint: the product/inventory data model, the seed catalog, the KPI set, the ML models, the AI prompts (voice/handwrite/invoice), the Vision shelf-scan and a lot of UI copy are grocery-coded. A store_type field already exists in the database but is dormant — it drives no behaviour today.")
    para("Recommendation in one line:", bold=True)
    para("Build a config-driven 'vertical pack' layer (one codebase, not per-vertical forks) plus a product variant/attribute model, then light up verticals in phases — Foundation, Apparel pilot, Electronics, then scale.")

    # 2. Scope
    heading("2. Scope", 1)
    para("In scope (general retail):", bold=True)
    bullets(VERTICALS_IN)
    para("Out of scope:", bold=True)
    bullets(VERTICALS_OUT)
    para("Pharmacy is excluded deliberately: drug licensing, scheduled-drug rules, prescription handling and mandatory batch+expiry make it a separate, heavily-regulated product rather than a 'vertical pack'.", italic=True)

    # 3. Current-state assessment
    heading("3. Current-state assessment", 1)
    para("Already generic — reuse as-is (strengths):", bold=True)
    bullets([
        "Auth & sessions, dual-token POS, subscription/Pro gating.",
        "Orders, checkout, payments, transaction history.",
        "Customers, segments, customer-specific pricing.",
        "Procurement & suppliers (needs light extension, not a rebuild).",
        "Baskets / bundles engine.",
        "Udhaar / credit ledger — used across Indian small retail (apparel, hardware, general), a cross-vertical strength, not a grocery-only feature.",
        "Localisation infrastructure (7 languages) and notification/FCM plumbing.",
    ])
    para("Grocery-coded — needs change:", bold=True)
    bullets([
        "Product/inventory model: fixed grocery attributes; no size x colour variant matrix; no serial/IMEI; no warranty.",
        "Expiry / perishable / batch is first-class and always-on.",
        "Units hardcoded (kg/g/L/loose); grocery example copy.",
        "KPIs, ML, AI prompts, Vision and WhatsApp are grocery-tuned.",
        "Tax/GST is parsed from invoices but never applied at billing.",
        "Branding & copy ('Kirana AI', Maida/Pulse examples, 'mykiranastore').",
    ])

    # 4. Strategic approach
    heading("4. Strategic approach — vertical packs", 1)
    para("Rather than fork the app per vertical (unmaintainable across 7 languages and two repos), introduce a single 'vertical' dimension that drives configuration. A Vertical Config selects, per store:")
    bullets([
        "Enabled features (expiry, loose/weight, variants, serial, warranty) via feature flags.",
        "Unit registry (apparel: piece/pair/set; electronics: unit; hardware: m/ft/pack).",
        "Attribute schema (size, colour, material, storage, shade, dimensions ...).",
        "KPI set (grocery vs apparel sell-through/size-curve vs electronics attach-rate/GMROI).",
        "ML profile (grocery daily-velocity vs apparel seasonal vs electronics slow/high-value).",
        "AI prompt pack (voice/handwrite/invoice grounded in the right vocabulary/catalog).",
        "Tax profile (HSN/category + price threshold -> GST rate).",
        "Copy pack (brand, examples, dashboard tone) across the 7 languages.",
    ])
    para("The keystone enabler is a product variant + dynamic attribute model so that 'one product, many sellable variants' works (size x colour for apparel, model/config for electronics), with grocery represented as a single implicit variant for full backward compatibility.")

    # 5. Gap analysis by theme (condensed table)
    heading("5. Gap analysis (by theme)", 1)
    para(f"{len(GAPS)} gaps are logged in the companion Excel workbook ('Gap Register' sheet) with layer, required change, affected verticals, priority and evidence. Summary by theme and priority:")
    # aggregate by theme
    from collections import OrderedDict
    theme_rows = OrderedDict()
    for g in GAPS:
        gid, theme, layer, cur, gap, change, verts, prio, eff, ev = g
        theme_rows.setdefault(theme, {"P0": 0, "P1": 0, "P2": 0, "ids": []})
        theme_rows[theme][prio] += 1
        theme_rows[theme]["ids"].append(gid)
    trows = [[t, str(d["P0"]), str(d["P1"]), str(d["P2"]), ", ".join(d["ids"])]
             for t, d in theme_rows.items()]
    table(["Theme", "P0", "P1", "P2", "Gap IDs"], trows, widths=[1.6, 0.5, 0.5, 0.5, 2.5])
    para("")
    para("The P0 (foundational) gaps that everything else depends on:", bold=True)
    bullets([f"{g[0]} — {g[5]}" for g in GAPS if g[7] == "P0"])

    # 6. Proposed data model
    heading("6. Proposed data-model changes", 1)
    para("New and modified tables (full columns in the Excel 'Proposed Data Model' sheet):")
    table(["Table", "Change", "Purpose"],
          [[d[0], d[1], d[2]] for d in DATAMODEL], widths=[1.8, 1.0, 4.0])

    # 7. Roadmap
    heading("7. Phased roadmap", 1)
    table(["Phase", "Timeline", "Outcome"],
          [[r[0], r[1], r[3]] for r in ROADMAP], widths=[2.6, 1.1, 3.0])
    para("")
    for ph in ROADMAP:
        para(f"{ph[0]}  ({ph[1]})", bold=True)
        para(ph[2])
        para(f"Gap IDs: {ph[4]}", italic=True, size=9)

    # 8. Risks
    heading("8. Risks & mitigations", 1)
    table(["Risk", "Mitigation"],
          [
            ["Variant/attribute migration could break live grocery stores & order history.",
             "Backward-compatible: grocery = single implicit variant; phase behind Vertical Config; verify on staging before any deploy."],
            ["ML cold-start for new verticals (no historical sales).",
             "Heuristic/rule-based recommendations first; collect data; train per-vertical models later."],
            ["Tax/GST correctness (compliance).",
             "Drive rates from an HSN/threshold table reviewed with finance; show breakup on bill; start with apparel/electronics brackets."],
            ["Catalog cold-start for non-grocery.",
             "Barcode lookup + fast manual add with attribute templates; ingest per-vertical catalogs over time."],
            ["Vision accuracy outside grocery.",
             "Gate Vision off where unsupported; redesign capture per vertical later; never block billing on it."],
            ["Scope creep into pharmacy.",
             "Keep pharmacy explicitly out; revisit only as a separate regulated product."],
            ["7-language localisation overhead for new copy.",
             "Parameterise copy packs; translate only the delta (proven workflow already in place)."],
          ], widths=[3.0, 3.6])

    # 9. Effort
    heading("9. Indicative effort & sequencing", 1)
    para("Rough, order-of-magnitude (2-4 engineers; refine after Phase 0 design):")
    bullets([
        "Phase 0 Foundation: ~4-6 weeks — unlocks General/Variety, Stationery, Hardware, Cosmetics with light extra work.",
        "Phase 1 Apparel pilot: ~6-8 weeks — first true non-grocery vertical, validates the model.",
        "Phase 2 Electronics: ~6-8 weeks — serial/IMEI, warranty, EMI, 18/28% GST.",
        "Phase 3 Scale & polish: ongoing — catalogs, vertical AI/WhatsApp/Vision, ML retraining, reporting.",
    ])
    para("First measurable milestone: a non-grocery store (e.g. a stationery or general store) fully operating on the Foundation layer, followed by the apparel pilot.", italic=True)

    # Appendix
    heading("Appendix — companion workbook", 1)
    para("Retail_Expansion_Gap_Analysis.xlsx contains: Executive Summary, the full Gap Register (filterable), the Feature x Vertical Matrix, the Phased Roadmap, and the Proposed Data Model. This document and the workbook are generated from docs/generate_gap_analysis.py and can be regenerated.")

    doc.save(path)
    return path


if __name__ == "__main__":
    xlsx = build_excel(os.path.join(OUT_DIR, "Retail_Expansion_Gap_Analysis.xlsx"))
    docx_path = build_word(os.path.join(OUT_DIR, "Retail_Expansion_Gap_Analysis.docx"))
    print("WROTE:", xlsx)
    print("WROTE:", docx_path)
