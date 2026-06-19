"""Convert the legal Markdown drafts in this folder to .docx (Word) files.

Handles the subset of Markdown used in these documents: H1-H3 headings,
blockquotes, pipe tables, bullet lists, horizontal rules, and inline
**bold**, *italic*, and `code` spans. Run:  python md_to_docx.py
"""
import glob
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\)|\[[^\]]+\])")


def add_runs(paragraph, text):
    """Render inline bold/italic/code/links into runs on a paragraph."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
        else:
            # Markdown link [text](url) -> "text (url)"; bare [PLACEHOLDER] kept.
            m = re.match(r"\[(.+?)\]\((.+?)\)", part)
            if m:
                paragraph.add_run(f"{m.group(1)} ({m.group(2)})")
            else:
                paragraph.add_run(part)


def flush_table(doc, rows):
    # rows: list of list-of-cell-strings; row[1] is the |---|---| separator.
    header = rows[0]
    body = rows[2:] if len(rows) > 1 else []
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, cell in enumerate(header):
        p = table.rows[0].cells[i].paragraphs[0]
        add_runs(p, cell)
        for r in p.runs:
            r.bold = True
    for row in body:
        cells = table.add_row().cells
        for i in range(len(header)):
            text = row[i] if i < len(row) else ""
            add_runs(cells[i].paragraphs[0], text)
    doc.add_paragraph()


def split_row(line):
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    table_buf = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Accumulate table rows.
        if line.startswith("|"):
            table_buf.append(split_row(line))
            i += 1
            continue
        elif table_buf:
            flush_table(doc, table_buf)
            table_buf = []

        if not line.strip():
            i += 1
            continue
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph().add_run("_" * 40).font.color.rgb = RGBColor(
                0xBB, 0xBB, 0xBB
            )
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, line.lstrip("> ").strip())
        elif re.match(r"\s*[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^\s*[-*] ", "", line))
        else:
            p = doc.add_paragraph()
            add_runs(p, line)
        i += 1

    if table_buf:
        flush_table(doc, table_buf)

    doc.save(docx_path)
    print(f"  -> {os.path.basename(docx_path)}")


def main():
    md_files = sorted(glob.glob(os.path.join(HERE, "*.md")))
    print(f"Converting {len(md_files)} Markdown file(s) to .docx:")
    for md in md_files:
        docx_path = os.path.splitext(md)[0] + ".docx"
        convert(md, docx_path)
    print("Done.")


if __name__ == "__main__":
    main()
